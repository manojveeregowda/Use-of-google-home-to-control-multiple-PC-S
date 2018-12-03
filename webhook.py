# ##################################### #
#   CSE 6349-001 Project -- Fall 2018   #
# ##################################### #
# Damian Jimenez             1000584863 #
# Imran Moiz Mohamad         1001535660 #
# Manoj Mysore Veere Gowda   1001547261 #
# ##################################### #

from flask import Flask, request
from flask_assistant import Assistant, ask, tell
from docx import Document, opc
from threading import Thread
from time import sleep
import json
import re
import logging
import subprocess
import os

# Reference: https://flask-assistant.readthedocs.io/en/latest/quick_start.html
logging.getLogger('flask_assistant').setLevel(logging.DEBUG)
app = Flask(__name__)
assist = Assistant(app, route='/')


@assist.action("greeting")
def greet_and_start():
    """
    Initial response to the user once the application starts or a greeting is detected.

    :return: Message greeting the user.
    """
    speech = "Hi! I'm Helper Holly, let me know how I can assist you."

    return ask(speech)


@assist.action("help")
def help_user():
    """
    Initial response to the user once the application starts or a greeting is detected.

    :return: Message letting the user know which commands are available and how to invoke them.
    """
    speech = "Having trouble? You can ask me to do various things like: opening a file, writing to a file, waking up" \
             " a computer, or putting a computer to sleep.\n\nHere are some examples: \"Open test.txt on John's " \
             "computer\", \"Type 'Hello World' in file notes.doc on Mary's computer\", \"Wake up Bill's " \
             "computer\", \"Put to sleep Jane's computer\"."

    return ask(speech)


@assist.action("quit")
def quit_app():
    """
    Quit the application

    :return: Message letting the user know the application is closing.
    """
    speech = "Okay, I hope I was of some help. Goodbye!"

    return tell(speech)


@assist.action("give-command", mapping={"given_name": "sys.given-name"})
def give_command(command, given_name):
    """
    Executes a command expressed by the user.

    :param command: command name
    :param given_name: the user's first name
    :return: Response that varies depending on the command that was executed
    """
    command_map[command]["function"](json.loads(request.get_data()))
    speech = command_map[command]["response"].format(given_name)

    return ask(speech)


def open_word(req=None):
    request_text = req["originalRequest"]["data"]["inputs"][0]["rawInputs"][0]["query"]
    try:
        filename = "{0}.docx".format(re.match(filename_re, request_text).group("filename"))
    except AttributeError:
        filename = ""

    command_list = req["target_computer_word_exe"]
    command_list.append(filename)

    subprocess.Popen(command_list)


def edit_word_file(req=None):
    request_text = req["originalRequest"]["data"]["inputs"][0]["rawInputs"][0]["query"]
    try:
        filename = "{0}.docx".format(re.match(filename_re, request_text).group("filename"))
    except AttributeError:
        filename = ""
    try:
        text = re.match(dictated_text_re, request_text).group("note")
        if "in file" in text:
            text = re.match(in_file_re, text).group("note")
    except AttributeError:
        text = None

    if text:
        if filename:
            try:
                doc = Document(filename)
            except opc.exceptions.PackageNotFoundError:
                doc = Document()
        else:
            doc = Document()
            filename = "demo.docx"

        doc.add_paragraph(text)
        doc.save(filename)


def sleep_computer(req=None):
    def sleep_helper():
        sleep(3)
        os.system("rundll32.exe powrprof.dll,SetSuspendState Sleep")

    t = Thread(target=sleep_helper)
    t.start()

    return None


def wake_up_computer(req=None):
    return None


if __name__ == '__main__':
    """
    Initialize the program and define the mapping of commands to actions and responses.
    
    Commands can control different PC's and are:
        - Open A Word Document
        - Type text into a word document
        - Putting a PC to sleep
        - Waking up a PC from sleep
    
    Windows Powershell commands to set environment variables:
        - $env:DEV_ACCESS_TOKEN = "your-dev-token-here"
        - $env:CLIENT_ACCESS_TOKEN = "your-client-token-here"
    
    Linux commands to set environment variables:
        - export DEV_ACCESS_TOKEN="your-dev-token-here"
        - export CLIENT_ACCESS_TOKEN="your-client-token-here"
    """
    filename_re = re.compile("(?:.*|^)(?:open|in) file (?P<filename>[a-zA-Z0-9.!@#$%^&*()\-_=+><\[\]{}|]+) .*$", re.IGNORECASE)
    dictated_text_re = re.compile("^(?:type|jot down|write) (?P<note>.*)", re.IGNORECASE)
    in_file_re = re.compile("^(?P<note>.*) in file.*$")
    command_map = {
        "open": {"function": open_word, "response": "Okay, I've opened Word for you."},
        "type": {"function": edit_word_file, "response": "Okay, I'm done writing what you asked."},
        "sleep": {"function": sleep_computer, "response": "Okay, putting {0}'s computer to sleep."},
        "wake_up": {"function": wake_up_computer, "response": "Okay, waking up {0}'s computer."}
    }

    app.run(debug=False, threaded=True)
